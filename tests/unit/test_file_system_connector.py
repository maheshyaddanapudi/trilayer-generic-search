from __future__ import annotations

import io
import sys
from pathlib import Path
from types import ModuleType
from unittest.mock import MagicMock, patch

import pytest

from src.connectors.file_system import (
    FileSystemConnector, RawSection, TextExtractor,
    _PDF_MIME, _DOCX_MIME, _XLSX_MIME, SUPPORTED_MIMES,
)


# ── Helpers for pdfminer sys.modules patching ─────────────────────────────────

class _FakeLTTextLine:
    def __init__(self, text: str) -> None:
        self._text = text

    def get_text(self) -> str:
        return self._text


class _FakeLTTextBox:
    def __init__(self, lines: list) -> None:
        self._lines = lines

    def __iter__(self):
        return iter(self._lines)


def _fake_pdfminer_modules(extract_pages_fn):
    """Return a context manager that injects fake pdfminer into sys.modules."""
    layout_mod = ModuleType("pdfminer.layout")
    layout_mod.LTTextBox = _FakeLTTextBox
    layout_mod.LTTextLine = _FakeLTTextLine

    highlevel_mod = ModuleType("pdfminer.high_level")
    highlevel_mod.extract_pages = extract_pages_fn

    pdfminer_mod = ModuleType("pdfminer")

    return patch.dict(sys.modules, {
        "pdfminer": pdfminer_mod,
        "pdfminer.high_level": highlevel_mod,
        "pdfminer.layout": layout_mod,
    })


# ── RawSection ────────────────────────────────────────────────────────────────

def test_raw_section_content_summary_truncates():
    s = RawSection(heading="H", content_text="x" * 500, page_number=1, order_index=0)
    assert len(s.content_summary) <= 300


def test_raw_section_short_content_unchanged():
    s = RawSection(heading="H", content_text="short", page_number=1, order_index=0)
    assert s.content_summary == "short"


# ── TextExtractor.extract_docx ────────────────────────────────────────────────

def _make_docx_bytes() -> bytes:
    from docx import Document
    doc = Document()
    doc.add_heading("Section One", level=1)
    doc.add_paragraph("Content of section one.")
    doc.add_heading("Section Two", level=2)
    doc.add_paragraph("Content of section two.")
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def test_extract_docx_sections():
    extractor = TextExtractor()
    sections = extractor.extract_docx(_make_docx_bytes())
    assert len(sections) == 2
    assert sections[0].heading == "Section One"
    assert "Content of section one" in sections[0].content_text


def test_extract_docx_no_headings():
    from docx import Document
    doc = Document()
    doc.add_paragraph("Just a paragraph.")
    buf = io.BytesIO()
    doc.save(buf)
    extractor = TextExtractor()
    sections = extractor.extract_docx(buf.getvalue())
    # No headings → one section with empty heading (Untitled)
    assert len(sections) == 1
    assert "Just a paragraph" in sections[0].content_text


# ── TextExtractor.extract_xlsx ────────────────────────────────────────────────

def _make_xlsx_bytes() -> bytes:
    from openpyxl import Workbook
    wb = Workbook()
    ws1 = wb.active
    ws1.title = "Sheet1"
    ws1.append(["Col A", "Col B"])
    ws1.append(["val1", "val2"])
    ws2 = wb.create_sheet("Sheet2")
    ws2.append(["data"])
    buf = io.BytesIO()
    wb.save(buf)
    return buf.getvalue()


def test_extract_xlsx_one_section_per_sheet():
    extractor = TextExtractor()
    sections = extractor.extract_xlsx(_make_xlsx_bytes())
    assert len(sections) == 2
    assert sections[0].heading == "Sheet1"
    assert sections[1].heading == "Sheet2"


def test_extract_xlsx_content_has_row_data():
    extractor = TextExtractor()
    sections = extractor.extract_xlsx(_make_xlsx_bytes())
    assert "Col A" in sections[0].content_text or "val1" in sections[0].content_text


# ── TextExtractor.extract_pdf (mocked pdfminer) ───────────────────────────────

def test_extract_pdf_with_sections():
    box = _FakeLTTextBox([
        _FakeLTTextLine("INTRODUCTION\n"),
        _FakeLTTextLine("Some content here.\n"),
    ])
    page = MagicMock()
    page.__iter__ = MagicMock(return_value=iter([box]))

    def fake_extract_pages(_io):
        return iter([page])

    with _fake_pdfminer_modules(fake_extract_pages):
        extractor = TextExtractor()
        sections = extractor.extract_pdf(b"fake pdf bytes")

    assert len(sections) >= 1
    assert sections[0].heading == "INTRODUCTION"


def test_extract_pdf_no_headings():
    box = _FakeLTTextBox([_FakeLTTextLine("ordinary text\n")])
    page = MagicMock()
    page.__iter__ = MagicMock(return_value=iter([box]))

    def fake_extract_pages(_io):
        return iter([page])

    with _fake_pdfminer_modules(fake_extract_pages):
        extractor = TextExtractor()
        sections = extractor.extract_pdf(b"fake")

    assert len(sections) == 1


def test_extract_pdf_empty():
    with _fake_pdfminer_modules(lambda _io: iter([])):
        extractor = TextExtractor()
        sections = extractor.extract_pdf(b"")
    assert sections == []


# ── TextExtractor._detect_sections_heuristic ──────────────────────────────────

def test_detect_sections_heuristic_returns_empty():
    extractor = TextExtractor()
    result = extractor._detect_sections_heuristic(["line 1", "line 2"])
    assert result == []


# ── FileSystemConnector ────────────────────────────────────────────────────────

def test_process_upload_pdf(tmp_path):
    conn = FileSystemConnector(watch_dir=tmp_path)
    mock_extractor = MagicMock()
    mock_extractor.extract_pdf.return_value = [
        RawSection("§1 Scope", "Scope content", 1, 0),
        RawSection("§2 Policy", "Policy content", 2, 1),
    ]
    conn._extractor = mock_extractor

    entities = list(conn.process_upload(b"pdf bytes", "policy.pdf", _PDF_MIME))
    doc = next(e for e in entities if e.entity_type == "Document")
    sections = [e for e in entities if e.entity_type == "Section"]
    assert doc.name == "policy"
    assert len(sections) == 2


def test_process_upload_docx(tmp_path):
    conn = FileSystemConnector(watch_dir=tmp_path)
    conn._extractor = MagicMock()
    conn._extractor.extract_docx.return_value = [
        RawSection("Heading 1", "content", 1, 0),
    ]
    entities = list(conn.process_upload(b"docx", "doc.docx", _DOCX_MIME))
    assert any(e.entity_type == "Section" for e in entities)


def test_process_upload_xlsx(tmp_path):
    conn = FileSystemConnector(watch_dir=tmp_path)
    conn._extractor = MagicMock()
    conn._extractor.extract_xlsx.return_value = [
        RawSection("Sheet1", "data", 1, 0),
    ]
    entities = list(conn.process_upload(b"xlsx", "data.xlsx", _XLSX_MIME))
    assert any(e.entity_type == "Section" for e in entities)


def test_process_upload_unsupported_mime(tmp_path):
    conn = FileSystemConnector(watch_dir=tmp_path)
    with pytest.raises(ValueError, match="Unsupported MIME"):
        list(conn.process_upload(b"data", "file.txt", "text/plain"))


def test_process_upload_zero_sections_fallback(tmp_path):
    conn = FileSystemConnector(watch_dir=tmp_path)
    conn._extractor = MagicMock()
    conn._extractor.extract_pdf.return_value = []
    entities = list(conn.process_upload(b"pdf", "empty.pdf", _PDF_MIME))
    sections = [e for e in entities if e.entity_type == "Section"]
    assert len(sections) == 1  # fallback single section


def test_section_entity_has_parent_relationship(tmp_path):
    conn = FileSystemConnector(watch_dir=tmp_path)
    conn._extractor = MagicMock()
    conn._extractor.extract_docx.return_value = [RawSection("H1", "content", 1, 0)]
    entities = list(conn.process_upload(b"d", "doc.docx", _DOCX_MIME))
    section = next(e for e in entities if e.entity_type == "Section")
    assert any(r.relation_type == "HAS_SECTION" for r in section.relationships)


def test_connect_creates_watch_dir(tmp_path):
    watch = tmp_path / "uploads"
    conn = FileSystemConnector(watch_dir=watch)
    conn.connect()
    assert watch.exists()


def test_disconnect_noop(tmp_path):
    conn = FileSystemConnector(watch_dir=tmp_path)
    conn.disconnect()  # no exception


def test_fetch_all_processes_files(tmp_path):
    # Write a .docx file
    from docx import Document
    doc = Document()
    doc.add_heading("H1", level=1)
    doc.add_paragraph("body")
    buf = io.BytesIO()
    doc.save(buf)
    (tmp_path / "test.docx").write_bytes(buf.getvalue())

    conn = FileSystemConnector(watch_dir=tmp_path)
    conn.connect()
    entities = list(conn.fetch_all())
    assert any(e.entity_type == "Document" for e in entities)


def test_fetch_all_skips_unknown_extension(tmp_path):
    (tmp_path / "file.csv").write_text("a,b,c")
    conn = FileSystemConnector(watch_dir=tmp_path)
    conn.connect()
    entities = list(conn.fetch_all())
    assert entities == []


def test_untitled_heading_gets_section_number(tmp_path):
    conn = FileSystemConnector(watch_dir=tmp_path)
    conn._extractor = MagicMock()
    conn._extractor.extract_pdf.return_value = [
        RawSection("Untitled", "content", 1, 0),
    ]
    entities = list(conn.process_upload(b"pdf", "my_doc.pdf", _PDF_MIME))
    section = next(e for e in entities if e.entity_type == "Section")
    assert "my_doc §1" in section.name


def test_supported_mimes_constant():
    assert _PDF_MIME in SUPPORTED_MIMES
    assert _DOCX_MIME in SUPPORTED_MIMES
    assert _XLSX_MIME in SUPPORTED_MIMES


# ── PDF edge cases — more coverage ───────────────────────────────────────────

def test_extract_pdf_non_textbox_elements_skipped():
    # L48: element that is NOT FakeLTTextBox → isinstance check fails → continue
    box = _FakeLTTextBox([_FakeLTTextLine("SECTION ONE\n")])
    non_box = "not a textbox"  # Not an LTTextBox instance

    page = MagicMock()
    page.__iter__ = MagicMock(return_value=iter([non_box, box]))

    def fake_extract_pages(_io):
        return iter([page])

    with _fake_pdfminer_modules(fake_extract_pages):
        extractor = TextExtractor()
        sections = extractor.extract_pdf(b"fake")

    assert len(sections) >= 1


def test_extract_pdf_non_textline_elements_skipped():
    # L51: element in LTTextBox that is NOT FakeLTTextLine → isinstance check fails → continue
    non_line = "not a text line"

    box = _FakeLTTextBox([non_line, _FakeLTTextLine("INTRO\n")])
    page = MagicMock()
    page.__iter__ = MagicMock(return_value=iter([box]))

    def fake_extract_pages(_io):
        return iter([page])

    with _fake_pdfminer_modules(fake_extract_pages):
        extractor = TextExtractor()
        sections = extractor.extract_pdf(b"fake")

    assert len(sections) >= 1


def test_extract_pdf_empty_lines_skipped():
    # L54: line with empty text → stripped = "" → continue
    box = _FakeLTTextBox([
        _FakeLTTextLine("\n"),       # empty after strip
        _FakeLTTextLine("   \n"),    # whitespace only
        _FakeLTTextLine("CONTENT\n"),
    ])
    page = MagicMock()
    page.__iter__ = MagicMock(return_value=iter([box]))

    def fake_extract_pages(_io):
        return iter([page])

    with _fake_pdfminer_modules(fake_extract_pages):
        extractor = TextExtractor()
        sections = extractor.extract_pdf(b"fake")

    # CONTENT is non-heading → goes to current_lines → final section with it
    assert len(sections) >= 1


def test_extract_pdf_two_headings_triggers_section_save():
    # L57: second heading found while first has content → section appended
    box = _FakeLTTextBox([
        _FakeLTTextLine("SECTION ONE\n"),   # heading 1
        _FakeLTTextLine("content here\n"),  # body
        _FakeLTTextLine("SECTION TWO\n"),   # heading 2 → L57 hit (saves section 1)
        _FakeLTTextLine("more content\n"),  # body
    ])
    page = MagicMock()
    page.__iter__ = MagicMock(return_value=iter([box]))

    def fake_extract_pages(_io):
        return iter([page])

    with _fake_pdfminer_modules(fake_extract_pages):
        extractor = TextExtractor()
        sections = extractor.extract_pdf(b"fake")

    assert len(sections) == 2
    assert sections[0].heading == "SECTION ONE"
    assert sections[1].heading == "SECTION TWO"


def test_extract_sections_unsupported_mime_returns_fallback(tmp_path):
    # L179: else branch in _extract_sections → sections=[] → fallback single section
    conn = FileSystemConnector(watch_dir=tmp_path)
    sections = conn._extract_sections(b"data", "application/unknown", "fallback_title")
    assert len(sections) == 1
    assert sections[0].heading == "fallback_title"


# ── Lucene _get_or_create opens existing index (L37) ─────────────────────────

def test_index_opens_existing_index_on_disk(tmp_path):
    from whoosh import index as whoosh_index
    from whoosh.fields import ID, TEXT, Schema

    _SCHEMA = Schema(
        chunk_id=ID(stored=True, unique=True),
        domain_id=ID(stored=True),
        entity_id=ID(stored=True),
        entity_type=ID(stored=True),
        breadcrumb=TEXT(stored=True),
    )
    from src.indexers.lucene import LuceneIndexWriter

    base = tmp_path / "whoosh_open_test"
    base.mkdir()
    domain_dir = base / "metadata"
    domain_dir.mkdir()

    # Create a valid Whoosh index synchronously (no AsyncWriter)
    ix = whoosh_index.create_in(str(domain_dir), _SCHEMA)
    ix.close()

    # Now a fresh writer calling index() will open the existing index (L37)
    writer = LuceneIndexWriter(index_dir=base)
    chunk = MagicMock()
    chunk.chunk_id = "c1"
    chunk.domain_id = "metadata"
    chunk.entity_id = "c1"
    chunk.entity_type = "Account"
    chunk.breadcrumb = "BC c1 revenue account"

    from src.models.core import MetadataChunk
    real_chunk = MetadataChunk(chunk_id="c1", domain_id="metadata", entity_id="c1",
                                entity_type="Account", breadcrumb="BC c1 revenue account")
    writer.index([real_chunk])  # Hits L37 (open_dir for existing index)
    results = writer.keyword_search("revenue", top_k=5, domain_id="metadata")
    assert any(r.chunk_id == "c1" for r in results)
